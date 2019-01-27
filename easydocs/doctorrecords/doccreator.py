from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from datetime import datetime
import os


def docCreate(patient, time, filename):
    PCF = Document()
    
    bday = datetime.strptime(patient.date_of_birth, '%Y/%m/%d')
    age = (datetime.today() - bday).days/365
    
    #set styles
    style = PCF.styles['Title']
    font1 = style.font
    font1.name = 'Arial'
    font1.size = Pt(30)

    style = PCF.styles['Heading 1']
    font2 = style.font
    font2.name = 'Arial'
    font2.size = Pt(18)
    font2.bold = False

    style = PCF.styles['Heading 2']
    font3 = style.font
    font3.name = 'Arial'
    font3.size = Pt(16)
    font3.bold = False

    style = PCF.styles['Normal']
    font4 = style.font
    font4.name = 'Arial'
    font4.size = Pt(12)

    #margins
    section = PCF.sections[0]
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)

    #header
    PCF.add_heading('Patient Consultation Form', 0)

    p1 = PCF.add_paragraph('')
    p1.add_run('Healthcare Practitioner: ').bold = True
    p1.add_run("\t\tDr. Melvin Hobbs")
    p1.style = PCF.styles['Normal']

    p = PCF.add_paragraph('')
    p.add_run('Date: \t\t\t\t\t').bold = True
    p.add_run(time.strftime("%Y-%m-%d"))

    PCF.add_heading('Patient Information', 1)
    PCF.add_paragraph('Name: \t\t' + str(patient.full_name()))
    PCF.add_paragraph('Gender: \t\t' + patient.sex)

    PCF.add_paragraph('Date of Birth: \t' + patient.date_of_birth.strftime("%Y-%m-%d") + '\t\t\t Age:\t' + age)
    PCF.add_paragraph('Address: \t\t' + str(patient.full_address()))

    PCF.add_paragraph('Phone: \t\t' + patient.phone_number)
    PCF.add_paragraph('HCN: \t\t\t' + patient.healthcard_number)

    PCF.add_heading('Family History', 1)
    PCF.add_paragraph(patient.family_history, style='List Bullet')
    p = PCF.add_paragraph('')
    p.add_run('Notes: ').bold = True

    PCF.add_heading("Active Medications", 1)
    for med in Medication
        PCF.add_heading(medication, 2)
        PCF.add_paragraph('Started: 				' +)
        PCF.add_paragraph('Perscription Ends:	' +)
        PCF.add_heading("Possible Side Effects")
        ###run loop for each side effect
    #   document.add_paragraph(
    #       'first item in unordered list', style='List Bullet'
    #   )
        PCF.add_paragraph('Notice Effects   Yes   No')
        p = PCF.add_paragraph('')
        p.add_run('Notes: ').bold = True

    PCF.add_heading('Ongoing Conditions', 1)
    ###run a loop for each Ongoing Condition
    #   PCF.add_heading(condition, 2)
    #   PCF.add_paragraph('Diagnosed: 				' +)
    #   PCF.add_paragraph('Method of Treatment:	' +)
    #   PCF.add_paragraph("New Treatments")
    #   PCF.add_paragraph("Incompatible Treatments:")
    #   p = PCF.add_paragraph('')
    #   p.add_run('Notes: ').bold = True

    PCF.add_heading('Potential Risks', 1)
    #   run a loop for each condition
    #   PCF.add_paragraph("Alzeimer's:")
    #   p1 = PCF.add_paragraph('\tAge, gender')

    PCF.add_heading('Additional Notes: ', 1)

    path_relative='export_docs/' + filename + '.docx'
    path = os.path.join(os.path.dirname(__file__), path_relative)
    
    print("Saving...")
    #create the form
    PCF.save(path)
