# Imports
from django.shortcuts import render
from django.http import HttpResponse
import zipfile

import datetime
from . import db_models as db
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
import os

















def zipdir(path, ziph):
    for root, dirs, files in os.walk(path):
        for file in files:
            ziph.write(os.path.join(root, file))

def pathify(path):
    if(path.endswith(".zip")):
        print("PATH" + os.path.dirname(__file__))
    return os.path.join(os.path.dirname(__file__), path)























def docCreate(appointment, time, filepath):
    patient = appointment.patient
    PCF = Document()
   
    #set styles
    style = PCF.styles['Title']
    font1 = style.font
    font1.name = 'Arial'
    font1.size = Pt(30)

    style = PCF.styles['Heading 1']
    font2 = style.font
    font2.name = 'Arial'
    font2.size = Pt(18)
    font2.bold = False

    style = PCF.styles['Heading 2']
    font3 = style.font
    font3.name = 'Arial'
    font3.size = Pt(16)
    font3.bold = False

    style = PCF.styles['Normal']
    font4 = style.font
    font4.name = 'Arial'
    font4.size = Pt(12)

    #margins
    section = PCF.sections[0]
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)

    #header
    PCF.add_heading('Patient Consultation Form', 0)

    p1 = PCF.add_paragraph('')
    p1.add_run('Healthcare Practitioner: ').bold = True
    p1.add_run("\tDr. Melvin Hobbs")
    p1.style = PCF.styles['Normal']

    p = PCF.add_paragraph('')
    p.add_run('Date of Appointment: \t\t').bold = True
    p.add_run(appointment.date.strftime("%Y-%m-%d") + "\t\t")

    p.add_run(appointment.strftime("%Y-%m-%d") + "\t\t")

    p.add_run('Date of Appointment: \t\t').bold = True
    p.add_run(time.strftime("%Y-%m-%d"))
    
    p = PCF.add_paragraph('')
    p.add_run('Time of Appointment: \t\t\t').bold = True
    p.add_run(appointment.time.strftime("%H:%M:%S"))

    PCF.add_heading('Patient Information', 1)
    PCF.add_paragraph('Name: \t\t' + str(patient.full_name()))
    PCF.add_paragraph('Gender: \t\t' + patient.sex)

    PCF.add_paragraph('Date of Birth: \t' + patient.date_of_birth.strftime("%Y-%m-%d") + '\t\t\t Age:\t' + str(patient.get_patient_age()))
    PCF.add_paragraph('Address: \t\t' + str(patient.full_address()))

    PCF.add_paragraph('Phone: \t\t' + patient.phone_number)
    PCF.add_paragraph('HCN: \t\t\t' + patient.healthcard_number)

    PCF.add_heading('Family History', 1)
    for fam in patient.family_history.all():
        PCF.add_paragraph(fam.family_member)
        for condition in fam.condition.all():
            PCF.add_paragraph(condition.name, style='List Bullet')
    p = PCF.add_paragraph('')
    p.add_run('Notes: ').bold = True
    PCF.add_paragraph()

    PCF.add_heading('Ongoing Conditions', 1)
    for OC in patient.active_conditions.all():
        PCF.add_heading(OC.condition.name, 2)
        PCF.add_paragraph('Diagnosed:\t\t\t\t\t' + OC.diagnosis_date.strftime("%Y-%m-%d"))
        PCF.add_paragraph('Method of Treatment:\t\t\t' + OC.get_treatment().name)
        PCF.add_paragraph('Started:\t\t\t\t\t' + OC.treatment_start_date.strftime("%Y-%m-%d"))
        PCF.add_paragraph('Perscription Ends:\t\t\t\t' + OC.treatment_renewal_date.strftime("%Y-%m-%d"))
        PCF.add_heading("Possible Side Effects")
        for SE in OC.condition.side_effects.all():
            PCF.add_paragraph(SE.effect, style='List Bullet')
        PCF.add_paragraph("Side Effects Noticed\tYes\tNo")
        p = PCF.add_paragraph("Recently Discovered Treatments:\t")
        for NT in OC.condition.new_treatments().all():    
            p.add_run(NT.name + ", ").italic = True
        PCF.add_paragraph("Incompatible Treatments:\t\t\t" + get_conflicting_meds(patient))
        p = PCF.add_paragraph('')
        p.add_run('Notes: ').bold = True
        PCF.add_paragraph()

    PCF.add_heading('Potential Risks', 1)
    for i in patient.is_patient_at_risk():
        PCF.add_paragraph(i, style='List Bullet')

    PCF.add_heading('Additional Notes: ', 1)
    
    #create the form
    PCF.save(filepath)






















# Create your views here.s
def homepage(request):
    hcpList = db.HealthcareProviders.objects.all()

    dt = datetime.datetime
    now = dt.now()

    hcp_selected = None
    appointment_list = None
    download_link = None

    date_start_str = request.POST.get('date_Start')
    date_end_str   = request.POST.get('date_End')

    date_start = dt.strptime(date_start_str, '%Y-%m-%d') if date_start_str else None
    date_end = dt.strptime(date_end_str, '%Y-%m-%d') if date_end_str else None
    
    if request.POST:

        if date_start and date_end and date_end < date_start:
            date_end = date_start

        if request.POST.get("btn_Today"):
            date_start = now
            date_end = now
            date_start_str = now.strftime("%Y-%m-%d")
            date_end_str = now.strftime("%Y-%m-%d")

        hcp_id = (int) (request.POST.get('dds_HealthcareProviders'))

        appointment_gen_list = request.POST.getlist('appointments')

        if appointment_gen_list:
            appointment_gen_list = list(filter(None, appointment_gen_list))
            appointment_gen_list = list(map(lambda x: db.Appointments.objects.get(pk=x), appointment_gen_list))
            appointment_gen_list = list(map(lambda x: [x, str(x) + ".docx"], appointment_gen_list))

        if request.POST.get('generate'):
            #delete existing files
            for filename in os.listdir(pathify('export_docs')):
                if not filename.startswith('.'):
                    os.unlink(pathify('export_docs' + '/' + filename))
            for filename in os.listdir(pathify('downloads')):
                if not filename.startswith('.'):
                    os.unlink(pathify('downloads' + '/' + filename))

            #generate new ones
            for apt in appointment_gen_list:
                docCreate(apt[0], now, pathify("export_docs/" +apt[1]))

            #zip?
            if len(appointment_gen_list) > 1:
                zipf = zipfile.ZipFile(pathify('downloads/docsdocs.zip'), 'w', zipfile.ZIP_STORED)
                zipdir(pathify('export_docs'), zipf)
                zipf.close()
                download_link = "/download/zip/docsdocs.zip"
            else:
                download_link = "/download/docx/" + appointment_gen_list[0][1]

        if (hcp_id != -1):
            hcp_selected = db.HealthcareProviders.objects.get(pk=hcp_id)

            if date_start and date_end: 
                appointment_list = db.Appointments.objects.filter(healthcare_provider=hcp_selected).filter(date__range=[date_start_str, date_end_str]).order_by('date')


    return render(request, 'home.html', {
        'HealthcareProviders':      hcpList,
        'SelectedHCP':              hcp_selected,
        'Appointments':             appointment_list,
        'StartDate':                date_start,
        'EndDate':                  date_end,
        'DownloadLink':             download_link
    })













def downloadzip(request, filename):
    with open(os.path.join(os.path.dirname(__file__), "downloads/" + filename) , 'rb') as myzipfile:
        response = HttpResponse(myzipfile.read())
        response['content_type'] = 'application/zip'
        response['Content-Disposition'] = 'attachment;filename=file.zip'
        return response

def downloaddocx(request, filename):
    with open(pathify("export_docs/" + filename) , 'rb') as docx:
        response = HttpResponse(docx.read())
        response['content_type'] = 'application/docx'
        response['Content-Disposition'] = 'attachment;filename=file.docx'
        return response

#def download(request, filename):
    
#CALL THIS FUNCTION ABOVE WITH A PATIENT
def get_conflicting_meds(patient):
    curr_treatment = patient.treatments
    incompat_treatment = db.Incompatibilities.objects.filter(treatment=curr_treatment)
    return incompat_treatment
    
